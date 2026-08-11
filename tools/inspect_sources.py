from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree
from openpyxl import load_workbook
from pypdf import PdfReader


NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def qname(local: str) -> str:
    return f"{{{NS['w']}}}{local}"


def text_of_paragraph(paragraph) -> str:
    return paragraph.text.replace("\t", "\\t").replace("\n", "\\n")


def inspect_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        paragraphs.append(
            {
                "index": index,
                "text": text_of_paragraph(paragraph),
                "style": paragraph.style.name if paragraph.style else None,
                "alignment": str(paragraph.alignment),
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": bool(run.underline) if run.underline is not None else None,
                        "font": run.font.name,
                        "size_pt": run.font.size.pt if run.font.size else None,
                    }
                    for run in paragraph.runs
                ],
            }
        )

    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            rows.append(
                {
                    "row": ri,
                    "height_twips": int(row.height.twips) if row.height else None,
                    "cells": [
                        {
                            "col": ci,
                            "text": cell.text.replace("\n", "\\n"),
                            "width_twips": int(cell.width.twips) if cell.width else None,
                            "paragraphs": [
                                {
                                    "text": text_of_paragraph(p),
                                    "style": p.style.name if p.style else None,
                                    "alignment": str(p.alignment),
                                    "runs": [
                                        {
                                            "text": r.text,
                                            "bold": r.bold,
                                            "font": r.font.name,
                                            "size_pt": r.font.size.pt if r.font.size else None,
                                        }
                                        for r in p.runs
                                    ],
                                }
                                for p in cell.paragraphs
                            ],
                        }
                        for ci, cell in enumerate(row.cells)
                    ],
                }
            )
        tables.append({"index": ti, "style": table.style.name if table.style else None, "rows": rows})

    sections = []
    for index, section in enumerate(doc.sections):
        sections.append(
            {
                "index": index,
                "page_width_twips": section.page_width.twips,
                "page_height_twips": section.page_height.twips,
                "orientation": str(section.orientation),
                "margins_twips": {
                    "top": section.top_margin.twips,
                    "right": section.right_margin.twips,
                    "bottom": section.bottom_margin.twips,
                    "left": section.left_margin.twips,
                    "header": section.header_distance.twips,
                    "footer": section.footer_distance.twips,
                },
                "header": [text_of_paragraph(p) for p in section.header.paragraphs],
                "footer": [text_of_paragraph(p) for p in section.footer.paragraphs],
            }
        )

    package = {}
    with zipfile.ZipFile(path) as archive:
        for info in archive.infolist():
            package[info.filename] = {"size": info.file_size, "crc": info.CRC}
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        textboxes = []
        for node in document_xml.xpath(".//w:txbxContent", namespaces=NS):
            text = "".join(node.xpath(".//w:t/text()", namespaces=NS))
            textboxes.append(text)
        drawings = len(document_xml.xpath(".//w:drawing", namespaces=NS))
        picts = len(document_xml.xpath(".//w:pict", namespaces=NS))
        bookmarks = [
            {"name": n.get(qname("name")), "id": n.get(qname("id"))}
            for n in document_xml.xpath(".//w:bookmarkStart", namespaces=NS)
        ]
        controls = []
        for control in document_xml.xpath(".//w:sdt", namespaces=NS):
            tag = control.find(".//w:tag", namespaces=NS)
            alias = control.find(".//w:alias", namespaces=NS)
            controls.append(
                {
                    "tag": tag.get(qname("val")) if tag is not None else None,
                    "alias": alias.get(qname("val")) if alias is not None else None,
                    "text": "".join(control.xpath(".//w:sdtContent//w:t/text()", namespaces=NS)),
                }
            )

    return {
        "type": "docx",
        "path": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
        "sections": sections,
        "inline_shapes": len(doc.inline_shapes),
        "drawings": drawings,
        "legacy_pictures": picts,
        "textboxes": textboxes,
        "bookmarks": bookmarks,
        "content_controls": controls,
        "package": package,
    }


def cell_summary(cell) -> dict:
    fill = cell.fill
    font = cell.font
    border = cell.border
    return {
        "coordinate": cell.coordinate,
        "value": cell.value,
        "data_type": cell.data_type,
        "number_format": cell.number_format,
        "style_id": cell.style_id,
        "font": {
            "name": font.name,
            "size": font.sz,
            "bold": font.b,
            "italic": font.i,
            "color_type": font.color.type if font.color else None,
            "color": font.color.rgb if font.color and font.color.type == "rgb" else None,
        },
        "fill": {
            "type": fill.fill_type,
            "fg_type": fill.fgColor.type,
            "fg": fill.fgColor.rgb if fill.fgColor.type == "rgb" else fill.fgColor.indexed,
        },
        "alignment": {
            "horizontal": cell.alignment.horizontal,
            "vertical": cell.alignment.vertical,
            "wrap": cell.alignment.wrap_text,
        },
        "border": {
            "left": border.left.style,
            "right": border.right.style,
            "top": border.top.style,
            "bottom": border.bottom.style,
        },
    }


def inspect_xlsx(path: Path) -> dict:
    workbook = load_workbook(path, data_only=False)
    sheets = []
    for sheet in workbook.worksheets:
        cells = []
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None or cell.style_id != 0:
                    cells.append(cell_summary(cell))
        sheets.append(
            {
                "title": sheet.title,
                "state": sheet.sheet_state,
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
                "freeze_panes": str(sheet.freeze_panes) if sheet.freeze_panes else None,
                "auto_filter": str(sheet.auto_filter.ref) if sheet.auto_filter.ref else None,
                "merged_cells": [str(rng) for rng in sheet.merged_cells.ranges],
                "column_dimensions": {
                    key: {"width": dim.width, "hidden": dim.hidden, "outline": dim.outlineLevel}
                    for key, dim in sheet.column_dimensions.items()
                },
                "row_dimensions": {
                    str(key): {"height": dim.height, "hidden": dim.hidden, "outline": dim.outlineLevel}
                    for key, dim in sheet.row_dimensions.items()
                    if dim.height is not None or dim.hidden or dim.outlineLevel
                },
                "print_area": str(sheet.print_area),
                "print_title_rows": sheet.print_title_rows,
                "page_setup": {
                    "orientation": sheet.page_setup.orientation,
                    "paper_size": sheet.page_setup.paperSize,
                    "fit_to_width": sheet.page_setup.fitToWidth,
                    "fit_to_height": sheet.page_setup.fitToHeight,
                },
                "cells": cells,
            }
        )
    return {
        "type": "xlsx",
        "path": str(path),
        "defined_names": [str(name) for name in workbook.defined_names.values()],
        "sheets": sheets,
    }


def inspect_pdf(path: Path) -> dict:
    reader = PdfReader(path)
    pages = []
    for index, page in enumerate(reader.pages):
        pages.append(
            {
                "index": index + 1,
                "width_pt": float(page.mediabox.width),
                "height_pt": float(page.mediabox.height),
                "rotation": page.rotation,
                "text": page.extract_text() or "",
            }
        )
    return {
        "type": "pdf",
        "path": str(path),
        "metadata": {str(k): str(v) for k, v in (reader.metadata or {}).items()},
        "pages": pages,
    }


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    root = Path(sys.argv[1] if len(sys.argv) > 1 else ".").resolve()
    out_dir = Path(sys.argv[2] if len(sys.argv) > 2 else root / "tmp" / "source-inspection")
    out_dir.mkdir(parents=True, exist_ok=True)
    manifest = []
    for path in sorted(root.iterdir()):
        if path.suffix.lower() not in {".docx", ".xlsx", ".pdf"}:
            continue
        if path.suffix.lower() == ".docx":
            report = inspect_docx(path)
        elif path.suffix.lower() == ".xlsx":
            report = inspect_xlsx(path)
        else:
            report = inspect_pdf(path)
        report_path = out_dir / f"{path.name}.json"
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
        manifest.append({"source": path.name, "report": str(report_path), "type": report["type"]})
        print(f"inspected\t{path.name}\t{report_path}")
    (out_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
